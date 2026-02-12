from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import json

def generate_radiology_report(data: dict, template_info: dict, output_path: str):
    """
    Generate a professional radiology report in DOCX format.
    data: Extracted and verified data from the user.
    template_info: The original JSON template with labels and sections.
    """
    doc = Document()
    
    # Header
    header = doc.add_heading(template_info.get("organ", "Radiology Report"), 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Template ID: {template_info.get('template_id', 'N/A')}")
    doc.add_paragraph(f"Protocol Posting Date: {template_info.get('protocol_posting_date', 'N/A')}")
    doc.add_paragraph("-" * 50)
    
    # Sections and Fields
    for section in template_info.get("sections", []):
        doc.add_heading(section.get("section_name", "SECTION"), level=1)
        
        for field in section.get("fields", []):
            field_id = field.get("field_id")
            label = field.get("label")
            value = data.get(field_id)
            
            p = doc.add_paragraph()
            p.add_run(f"{label}: ").bold = True
            
            display_value = "not determined"
            if value and value != "not_determined":
                # Try to map value to label if options exist
                if "options" in field:
                    option_label = next((opt["label"] for opt in field["options"] if opt["value"] == value), None)
                    display_value = option_label if option_label else value
                else:
                    display_value = value
            
            run = p.add_run(str(display_value))
            if display_value == "not determined":
                run.bold = True
    
    # Save the document
    doc.save(output_path)
    return str(output_path)

if __name__ == "__main__":
    # Test generation
    sample_data = {"procedure": "adrenalectomy_total", "specimen_laterality": "right"}
    with open("CAP templates/JSON_Output/Adrenal_4.3.1.0.REL_CAPCP.json", "r") as f:
        template = json.load(f)
    generate_radiology_report(sample_data, template, "test_report.docx")
